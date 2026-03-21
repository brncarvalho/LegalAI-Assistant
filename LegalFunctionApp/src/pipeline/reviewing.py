"""
Core legal review pipeline.

Contains functions for clause review, filtering, document generation,
deduplication, and search index operations.
"""

import json
from src.config.load_config import get_model_config
from src.pipeline.clause_extraction_and_processing import normalize_clause_number
from src.pipeline.embedding import generate_embedding
from azure.search.documents.models import VectorizableTextQuery
from azure.core.exceptions import HttpResponseError
import time
import tiktoken
import os
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
import zipfile
import shutil
from lxml import etree
import uuid
from azure.search.documents.indexes.models import (
    SearchField,
    SearchFieldDataType,
    SearchIndex,
    VectorSearch,
    HnswAlgorithmConfiguration,
    VectorSearchProfile,
    AzureOpenAIVectorizer,
    AzureOpenAIVectorizerParameters,
    SearchIndexerDataNoneIdentity,
)


def review_clauses(clauses, client, search_client, response_format, config, termo):
    """
    Review and legally validate a list of contract clauses using Azure Cognitive Search
    for reference retrieval and Azure OpenAI for structured legal review.

    Parameters:
        clauses (List[dict]): Each dict must have 'content' and 'clause_number' keys.
        client (AzureOpenAI): Initialized OpenAI client for chat completions.
        response_format: The structured output format to request from the LLM.
        config (dict): Model configuration dict (e.g., max_tokens, temperature, top_p, deployment).

    Returns:
        dict: {
            "reviewed_clauses": Dict mapping each clause_number to its structured review result,
            "usage": {
                "prompt": total prompt tokens used,
                "completion": total completion tokens used,
                "total": total tokens used
            }
        }
    """

    # override the module-level config with the passed-in configuration
    model_config = config

    # get the tokenizer for counting and slicing tokens
    encoding = tiktoken.get_encoding("cl100k_base")
    # prepare the output container for all reviewed clauses
    list_of_reviewed_clauses = {}

    # initialize token usage counters
    total_prompt = 0
    total_completion = 0
    total_tokens = 0

    # iterate over each clause dict
    for clause in clauses:
        # extract raw clause text and its identifier
        clause_content = clause["content"]
        clause_number = clause["clause_number"]

        # skip empty content or missing clause number
        if not clause_content or not clause_number:
            print(f"Cláusula vazia encontrada: {clause['clause_number']}. Pulando...")
            continue

        # trim whitespace from content and number
        clause_content = clause_content.strip()
        clause_number = clause_number.strip()

        # encode the clause text to tokens to enforce the model limit
        tokens = encoding.encode(clause_content)
        if len(tokens) > 8191:
            print(
                f" Cláusula excedeu limite de tokens ({len(tokens)} tokens). Pulando..."
            )
            continue

        # prepare to collect top-k similar clauses from Azure Search
        retrieved_clauses = []

        # build a vector query for the clause content
        vector_query = VectorizableTextQuery(
            text=clause["content"],  # text to vectorize
            k_nearest_neighbors=5,  # number of neighbors to retrieve
            fields="text_vector",  # vector field in the index
        )

        # execute the search with vector query
        results = list(
            search_client.search(
                search_text=None,
                vector_queries=[vector_query],
                select=["chunk"],  # retrieve only the 'chunk' field
                top=5,  # limit to 5 results
            )
        )

        # collect the 'chunk' text from each search result
        for result in results:
            retrieved_clauses.append(result["chunk"])

        # join the retrieved reference clauses into one string
        reference_clauses = "\n\n".join(retrieved_clauses)

        # craft the user prompt with the clause and its legal references
        prompt = f"""
        Você é um **advogado sênior** responsável por revisar cláusulas contratuais, utilizando exclusivamente um conjunto de **cláusulas de referência** previamente aprovadas e assinadas pelo Departamento Jurídico da empresa.

        Embora você deva sempre fazer a análise mais correta e justa, você pode se enviesar levemente caso a Logicalis seja a **contratante** e se posicionar como tal para garantir a máxima proteção jurídica para a empresa, porém caso a Logicalis seja a **contratada** 
        você deve buscar o mínimo de responsabilidade para a empresa, garantindo assim a convergência entre o rigor da sua análise e os interesses da empresa.
        ---

        ## Instruções

        1. **Base normativa**  
        - Utilize apenas o conteúdo jurídico contido nas cláusulas de referência.  
        - Não introduza disposições que não constem em pelo menos uma referência.  
        - Quando as referências divergir em grau de proteção, adote a solução mais restritiva, desde que nenhuma referência a contrarie.

        2. **Critérios de avaliação**  
        - **Alinhada** A cláusula alvo reproduz o conteúdo jurídico essencial das referências. → Mantenha‑a sem alterações.  
        - **Parcialmente alinhada** Existe lacuna pontual (ex.: ausência de exceção legal, de limite de responsabilidade, de prazo). → Inserir apenas o ponto faltante.  
        - **Desalinhada** A cláusula conflita com o padrão (ex.: amplia responsabilidade; omite obrigação mandatória). → Ajustar minimamente para restabelecer a conformidade.

        3. **Regras e princípios de redação**  

        3.1 **Mudança mínima** Preservar numeração, termos definidos, capitalização, estrutura (listas, alíneas) e tom do contrato do cliente. Alterar somente o indispensável.  
        3.2 **Placeholders** Marcas de anonimização entre colchetes nas referências (ex.: `[PRAZO]`, `[NOME]`) **não são requisitos**. Ignore‑as como variáveis. Não insira esses placeholders na cláusula revisada, são somente máscaras de anonimização do banco, o foco é sempre o contéudo jurídico.  
        3.3 **Cláusulas que citam anexos ou documentos externos** :
            ‑ **Não** incorporar listas de anexos, números, versões ou descrições que venham apenas das referências.  
            ‑ Se faltar detalhamento ou enumeração de anexos, registre o fato em `problema_juridico`, mas **mantenha a cláusula original** — as Partes precisam revisar manualmente, pois o você não tem acesso aos anexos.  
            ‑ Só altere se houver conflito jurídico expresso (por exemplo, cláusula original concede prioridade ao anexo em detrimento do contrato, mas as referências impõem o inverso). 

             Exemplo de saída desejada:

             Clausula orignal:

             X. Constituem parte integrante e indissociável do presente Contrato os seus Anexos, conforme indicado na Lista de Anexos do Termo de Contratação ("Anexos").

             Clasula revisada:

             X. Constituem parte integrante e indissociável do presente Contrato os seus Anexos, conforme indicado na Lista de Anexos do Termo de Contratação ("Anexos").
             
             problema_juridico: A cláusula original não especifica os anexos que integram o contrato, nem estabelece regras para resolução de conflitos entre disposições do contrato e dos anexos, o que pode gerar ambiguidades. Como não tenho detalhamento dos
             anexos, mantive a cláusula sem alteração porém com as pontuações anteriores.
        
        
        3.4 **Sem correção gramatical** (a menos que a alteração jurídica exija).  
        3.5 **Sem comparações literais**; descreva o risco jurídico de forma objetiva.

        **Exemplo de Revisão**:

            Clausula original:

            X. Propriedade. Toda a informação Confidencial, a não ser que de outro modo tenha sido estabelecido por escrito entre
            as Partes, permanecerá sendo de propriedade da Parte que transmitir a Informação Confidencial, somente podendo ser
            usada pela parte receptora para os fins deste Acordo de Confidencialidade. Tais Informações Confidenciais, incluídas as
            cópias realizadas, em prazo razoável serão retornadas para a parte que as transmitiu, ou então destruídas pela parte receptora, tão logo solicitado pela parte transmissora e, em
            qualquer caso, na hipótese de término deste Acordo de Confidencialidade. A pedido da parte que transmitir a Informação Confidencial, a parte receptora deverá prontamente
            emitir uma declaração a ser assinada por seu representante legal, confirmando que toda a Informação Confidencial não retornada para a parte transmissora foi inteiramente destruída.

            Cláusula revisada:

             5 - Toda a informação Confidencial, a não ser que de outro modo tenha sido estabelecido por escrito entre as Partes, permanecerá sendo de propriedade da Parte que transmitir a Informação Confidencial, somente podendo ser usada pela parte receptora para os fins deste Acordo de Confidencialidade. Tais Informações Confidenciais, incluídas as cópias realizadas, em prazo razoável serão retornadas para a parte que as transmitiu, 
             ou então destruídas pela parte receptora, tão logo solicitado pela parte transmissora, exceto aquelas que sejam exigidas por lei ou regulamento para retenção pela parte receptora, ou que estejam armazenadas automaticamente em sistemas de backup acessíveis apenas em situações excepcionais e não no curso normal dos negócios. A pedido da parte que transmitir a Informação Confidencial, a parte receptora deverá prontamente emitir uma declaração a ser assinada por seu representante legal, 
             confirmando que toda a Informação Confidencial não retornada para a parte transmissora foi inteiramente destruída, ou que está armazenada exclusivamente conforme exceções legais ou técnicas acima mencionadas.  
            
             Justificativa da alteração (problema_jurídico): "A cláusula não aborda a possibilidade de retenção de informações confidenciais por motivos legais, nem detalha exceções ou limitações à destruição ou devolução de informações confidenciais, o que pode gerar conflitos com obrigações legais ou regulatórias."


        4. **Saída obrigatória (JSON)**  
        Devolva exatamente no formato abaixo:  

        ```json
        [
            "numero_da_clausula": "<identificador original>",
            "clasula_original": "<texto original>",
            "problema_juridico": "<descrição objetiva ou 'Cláusula alinhada juridicamente às referências.'>",
            "clausula_revisada": "<texto após ajustes (ou original, se nenhum ajuste)>"
        ] 

        
        A cláusula a seguir a ser revisada, a Logicalis está atuando como **{termo}**:
        
        📄 **Cláusula:**
        {clause["clause_number"] + clause["content"]}

        **Cláusulas de Referência (juridicamente aprovadas):**
        {reference_clauses}
        """

        # call the OpenAI chat completion with structured output parsing

        messages = [
            {"role": "system", "content": "Você é um revisor jurídico experiente."},
            {"role": "user", "content": prompt},
        ]

        response = safe_parse_with_retry(client, messages, response_format, config)

        # accumulate token usage from this call
        total_prompt += response.usage.prompt_tokens
        total_completion += response.usage.completion_tokens
        total_tokens += response.usage.total_tokens

        # extract the structured output object from the response
        structured_output = response.choices[0].message.parsed

        # serialize and store the parsed JSON under the clause number
        list_of_reviewed_clauses[clause["clause_number"]] = json.loads(
            structured_output.model_dump_json(indent=2)
        )

    # return all reviewed clauses and the aggregate usage metrics
    return {
        "reviewed_clauses": list_of_reviewed_clauses,
        "usage": {
            "prompt": total_prompt,
            "completion": total_completion,
            "total": total_tokens,
        },
    }


def review_clauses_with_contract_context(
    clauses, client, index_client, response_format, config
):
    """
    Review clauses for redundancy within the context of a previously indexed contract.

    This function iterates over each clause reviewed by a first assistant, retrieves
    semantically similar clauses from an Azure Cognitive Search index, and then asks
    an LLM to determine if the revision introduced redundancy. If redundancy is detected,
    it reverts the clause revision and updates the legal problem field accordingly.
    Finally, it merges the updated clause back into the search index.

    Parameters:
        clauses (dict): Mapping of page keys to dicts containing 'clauses' lists of
                        dicts with keys: 'numero_da_clausula', 'clasula_original',
                        'problema_juridico', 'clausula_revisada'.
        client (AzureOpenAI): OpenAI client for chat completions.
        index_client (SearchClient): Azure Cognitive Search client for document retrieval
                                     and merging.
        response_format: Format specifier for structured LLM output.
        config (dict): Model configuration (max_tokens, temperature, top_p, deployment).

    Returns:
        dict: {
            "reviewed_clauses": Dict mapping cada numero_da_clausula to its LLM-reviewed result,
            "usage": {
                "prompt": total prompt tokens used,
                "completion": total completion tokens used,
                "total": total tokens used
            }
        }
    """

    # use passed-in config for model parameters
    model_config = config

    # initialize container for all reviewed clauses
    list_of_reviewed_clauses = {}

    # counters for aggregate token usage
    total_prompt = 0
    total_completion = 0
    total_tokens = 0

    # iterate over each page in the clauses dict
    for page_key, page in clauses.items():
        # iterate through each clause on the page
        for clause in page["clauses"]:
            # extract relevant fields from the clause
            numero_da_clausula = clause["numero_da_clausula"]
            clasula_original = clause["clasula_original"]
            problema_juridico = clause["problema_juridico"]
            clausula_revisada = clause["clausula_revisada"]

            # prepare list to collect similar clauses from the index
            retrieved_clauses = []

            # build a vector query for the original clause text
            vector_query = VectorizableTextQuery(
                text=clasula_original,  # text to vector-search
                k_nearest_neighbors=5,  # retrieve top 5 neighbors
                fields="embedding",  # field in index storing embeddings
            )

            # execute the search query against the index
            results = list(
                index_client.search(
                    search_text=None,
                    vector_queries=[vector_query],
                    select=["*"],  # retrieve all fields
                    top=3,  # limit to 3 results
                )
            )

            # collect id and clause fields from each search hit
            for result in results:
                retrieved_clauses.append(
                    {
                        "id": result["id"],
                        "numero_da_clausula": result["numero_da_clausula"],
                        "clasula_original": result["clasula_original"],
                        "problema_juridico": result["problema_juridico"],
                        "clausula_revisada": result["clausula_revisada"],
                    }
                )

            # construct the prompt to detect redundancy based on retrieved clauses

            prompt = f"""
                Você é um especialista jurídico responsável por validar se a revisão de uma cláusula contratual está redundante ao compará-la com uma outra cláusula que contém o mesmo risco jurídico .
                
                EXEMPLO 1: REDUNDÂNCIA:

                Cláusula 1
                clausula_original  = "… deve ser devolvida ou destruída pela Receptora ao término do contrato …"  
                problema_juridico  = "Não define momento exato da devolução/destruição."  
                clausula_revisada  = "… será devolvida ou destruída pela Receptora, desde que identificada previamente …"

                Cláusula 2 (similar)  
                clausula_original  = "… serão retornadas ou destruídas pela parte receptora, tão logo solicitado …"  
                problema_juridico  = "Já contém obrigação de devolução/destruição."  
                clausula_revisada  = "… serão retornadas ou destruídas pela parte receptora ao término do contrato …"

                ✅ Diagnóstico: MESMO risco jurídico → redundância. 
                
                ✅ Ação: manter revisão apenas na cláusula de número MAIOR (2):

                    Cláusula 2 (similar)  
                    clausula_original  = "… serão retornadas ou destruídas pela parte receptora, tão logo solicitado …"  
                    problema_juridico  = "Já contém obrigação de devolução/destruição."  
                    clausula_revisada  = "… serão retornadas ou destruídas pela parte receptora ao término do contrato …"

                    Cláusula 1
                    clausula_original  = "… deve ser devolvida ou destruída pela Receptora ao término do contrato …"  
                    problema_juridico  = "Cláusula 1 possui o mesmo risco da Cláusula 2; será tratado na 2."    
                    clausula_revisada  = "… deve ser devolvida ou destruída pela Receptora ao término do contrato …"

                EXEMPLO 2 (não-redundante em ciclo):
                
                Cláusula X (#6.1.4) → já diz:  
                problema_juridico = "Cláusula 6.1.4 possui o mesmo risco da 6.1.5; será tratado na 6.1.5."  
                Cláusula Y (#6.1.5) (atual)  
                *Não revertida* porque X aponta para ela.  
                Portanto **NÃO** marcar redundância outra vez.  
                Resultado para 6.1.5: campos permanecem como estão.
                
                ──────────────────────── SUA TAREFA
                Compare a cláusula alvo {numero_da_clausula} com a similar {retrieved_clauses[1]["numero_da_clausula"]} através dos seguintes passos:

                1 - Observe os campos informações da cláusula {numero_da_clausula} e {retrieved_clauses[1]["numero_da_clausula"]}.
                 
                2 - Possuem o mesmo risco jurídico?* 
                    Se o número da clásula {numero_da_clausula} é **MENOR** que {retrieved_clauses[1]["numero_da_clausula"]}, devolva {numero_da_clausula} com os 3 campos
                    com alterações conforme especificado no exemplo "1":
                        clausula_original: inalterado
                        problema_juridico: "Cláusula {numero_da_clausula} possui o mesmo risco da Cláusula {retrieved_clauses[1]["numero_da_clausula"]}; será tratado na 2{retrieved_clauses[1]["numero_da_clausula"]}." 
                        clausula_revisada: clausula_original 
                    Se o número da clásula {numero_da_clausula} é **MAIOR** que {retrieved_clauses[1]["numero_da_clausula"]}, devolva {numero_da_clausula} com os 3 campos sem nenhuma alteração conforme especificado no exemplo "1":
                        clausula_original: inalterado
                        problema_juridico: inalterado
                        clausula_revisada: inalterado

                    Se no campo "problema_juridico" da cláusula {retrieved_clauses[1]["numero_da_clausula"]}, já estiver "Cláusula {numero_da_clausula} possui o mesmo risco da Cláusula {retrieved_clauses[1]["numero_da_clausula"]}; será tratado na
                    {retrieved_clauses[1]["numero_da_clausula"]}, devolva {numero_da_clausula} com os 3 campos sem nenhuma alteração conforme especificado no exemplo "2":
                        clausula_original: inalterado
                        problema_juridico: inalterado
                        clausula_revisada: inalterado

                3 - Se {numero_da_clausula} não possui nenhuma redundância com {retrieved_clauses[1]["numero_da_clausula"]}, devolva {numero_da_clausula} com os campos "clausula_original", "problema_juridico" e "clausula_revisada" completamente inalterados, do jeito que estão.

                REGRAS:
                
                1 - Você como especialista deve fazer toda uma análise de ambas as clasulas antes de inferir se houve ou não redundância na revisão. A cláusula similar está vindo do banco por similaridade semantica,
                isso não necessariamente siginifica que o risco jurídico é o mesmo.

                2 - Não se esqueça da do exemplo 2 e atente-se ao efeito "circular". Ou seja, se {numero_da_clausula} for redundante e vai ser tratado na {retrieved_clauses[1]["numero_da_clausula"]}, ou vice-versa, quando voce estiver olhando a {retrieved_clauses[1]["numero_da_clausula"]},
            vai estar especiificado que {numero_da_clausula} seria tratado na {retrieved_clauses[1]["numero_da_clausula"]}. Nesse cenário, {retrieved_clauses[1]["numero_da_clausula"]} permancece INALTERADO e não deve apontar de volta para {numero_da_clausula}. Se voce ignorar isso, comprometerá a revisão.
            
                3 - "Inalterado", significa devolver os campos exatamente como eles vieram.


                📄 **Informações da cláusula {numero_da_clausula}:**
                    clasula_original = {clasula_original}
                    problema_juridico= {problema_juridico}
                    clausula_revisada = {clausula_revisada}

                **Cláusula similar que provavelmente contém o mesmo risco jurídico: **
                    Clásula: {retrieved_clauses[1]["numero_da_clausula"]} \n
                    Clasula Original: {retrieved_clauses[1]["clasula_original"]} \n
                    Problema Juridico: {retrieved_clauses[1]["problema_juridico"]} \n
                    Clausula Revisada: {retrieved_clauses[1]["clausula_revisada"]} \n

            """

            # call the LLM with structured parsing
            response = client.beta.chat.completions.parse(
                messages=[
                    {
                        "role": "system",
                        "content": "Você é um revisor jurídico experiente.",
                    },
                    {"role": "user", "content": prompt},
                ],
                response_format=response_format,
                max_tokens=model_config["max_tokens"],
                temperature=model_config["temperature"],
                top_p=model_config["top_p"],
                model=model_config["deployment"],
            )

            # update token usage counters
            total_prompt += response.usage.prompt_tokens
            total_completion += response.usage.completion_tokens
            total_tokens += response.usage.total_tokens

            # parse the structured output from the LLM
            structured_output = response.choices[0].message.parsed

            # store the parsed review under the clause number
            list_of_reviewed_clauses[clause["numero_da_clausula"]] = json.loads(
                structured_output.model_dump_json(indent=2)
            )

            # merge the updated clause back into the search index
            index_client.merge_documents(
                documents=[
                    {
                        "id": retrieved_clauses[0]["id"],
                        "clasula_original": list_of_reviewed_clauses[
                            clause["numero_da_clausula"]
                        ]["clauses"][0]["clasula_original"],
                        "problema_juridico": list_of_reviewed_clauses[
                            clause["numero_da_clausula"]
                        ]["clauses"][0]["problema_juridico"],
                        "clausula_revisada": list_of_reviewed_clauses[
                            clause["numero_da_clausula"]
                        ]["clauses"][0]["clausula_revisada"],
                    }
                ]
            )

    # return the collection of reviewed clauses and usage statistics
    return {
        "reviewed_clauses": list_of_reviewed_clauses,
        "usage": {
            "prompt": total_prompt,
            "completion": total_completion,
            "total": total_tokens,
        },
    }


def filter_clauses_with_gpt4o(chunks, client, response_format, config):
    """
    Use an LLM to extract and filter contract clauses, sub-clauses, and items from text chunks.

    This function sends each page chunk to a GPT-4o model with a detailed prompt
    instructing it to identify clause structures (e.g. “3.”, “3.5.”, “a.”, “i.”),
    ignore headers, preserve numbering, and separate sub-items even if orphaned.
    The model’s structured JSON output is collected per page.

    Parameters:
        chunks (List[dict]): Each dict must contain 'page_number' and 'content' keys.
        client (AzureOpenAI): Initialized GPT-4o OpenAI client for chat completions.
        response_format: Desired structured output format for the LLM response.
        config (dict): Model parameters (max_tokens, temperature, top_p, deployment).

    Returns:
        dict: {
            "clauses": {page_number: filtered PageOutput dict, ...},
            "usage": {"prompt": int, "completion": int, "total": int}
        }
    """

    # use the provided config for model settings
    model_config = config

    # counters to accumulate token usage across all calls
    total_prompt = 0
    total_completion = 0
    total_tokens = 0

    # container for the final filtered clauses per page
    filtered_clauses = {}

    i = 0
    # iterate through each text chunk representing a page
    for chunk in chunks:
        # build the system prompt with detailed extraction rules
        prompt = """
        
            Você é um doutor especializado em análise de documentos jurídicos empresariais.

            ## Objetivo
            Extrair **somente cláusulas numeradas** (por exemplo, `3.`, `3.5.`, `4.1.2.`) e **agrupar dentro do conteúdo da cláusula** todos os seus **subitens** (por exemplo, `a)`, `b)`, `i)`, `ii)`).  
            **Não** separar subitens em itens independentes: eles devem aparecer **dentro do `content` da cláusula-pai**.

            ## Instruções de extração

            1. **Identificação de marcadores principais**  
            - Um marcador principal é qualquer numeração no formato decimal pontuado: `N.`, `N.N.`, `N.N.N.` (o ponto final pode ou não aparecer no texto original).  
                Exemplos válidos: `11.`, `12`, `28.2`, `4.1.2.`  
            - Se o número vier por extenso ou em algarismos romanos (ex.: **“Cláusula Décima Quarta”**, **“XII”**), **converta** para a forma decimal pontuada:  
                - “Cláusula Décima Quarta” → `14.`  
                - “XII” → `12.`

                **Exemplo (entrada):**
                    11.1. A CONTRATADA deverá manter os registros...
                    12. A Vigência do contrato se manterá...
                    ...
                    XII – As penalidades serão aplicadas..
                **Exemplo (saída):**
                ```json
                    [
                    "clause_number": "11.1", "content": "A CONTRATADA deverá manter os registros..." ,
                    "clause_number": "12", "content": "As penalidades serão aplicadas",
                    "clause_number": "14", "content": "Vigência" 
                    ]

            1A. **Proibição de sintetizar cláusula-pai**

               - Nunca crie um número de cláusula pai que não apareça explicitamente no texto disponível (página atual + até 3 seguintes).

               - Se existirem 2.1, 2.2, 2.3, mas não existir 2 visível, não produza a cláusula 2. Extraia cada 2.X separadamente.

               **Exemplo (entrada):**
                   "2.1. Constituem parte integrante...
                    2.2. Também integram este Contrato...
                    2.3. A partir da assinatura deste...
                    2.4. Em caso de divergência...
                    2.5. Eventuais obrigações complementares..."

                **Exemplo (saída):**
                    [
                     "clause_number": "2.1", "content": "Constituem parte integrante..." ,
                     "clause_number": "2.2", "content": "Também integram este Contrato...", 
                     "clause_number": "2.3", "content": "A partir da assinatura deste..." ,
                     "clause_number": "2.4", "content": "Em caso de divergência..." ,
                     "clause_number": "2.5", "content": "Eventuais obrigações complementares..." 
                    ]
                **Exemplo (saída incorreta — proibida):**
                [
            
                 "clause_number": "2",
                 "content": "Constituem parte integrante... Também integram este Contrato... A partir da assinatura... Em caso de divergência... Eventuais obrigações..."      
                ]

                

            2. **Delimitação do conteúdo da cláusula**  
            Para cada marcador principal `N` localizado:
            - Defina `clause_number = "N"` (padronize como `N`, `N.N` ou `N.N.N`, sem duplicar pontos finais).
            - O `content` é **todo o texto** a partir desse marcador até **antes** do próximo marcador principal subsequente.
            - Inclua no `content`:
                - O corpo textual imediatamente após `N`.
                - **Todos os subitens pertencentes a `N`**, como listas `a)`, `b)`, `i)`, `ii)`, `1)`, `2)`, **desde que apareçam antes do próximo marcador principal**.

            **Exemplo (entrada):**
                1.2. A CONTRATADA executará os serviços conforme as especificações técnicas:
                    a) Escopo técnico.
                    b) Prazos de execução.
                    c) Critérios de aceite.
                1.3. O CONTRATANTE deverá...
            
            **Exemplo (saída):**
                ```json
                [
                "clause_number": "1.2",
                "content": "A CONTRATADA executará os serviços conforme as especificações técnicas:\na) Escopo técnico.\nb) Prazos de execução.\nc) Critérios de aceite."
                ]

            3. **Subitens**  
            - Subitens **não** geram registros separados. Devem ficar **concatenados no `content` da cláusula-pai**, preservando a ordem e a marcação literal (`a)`, `b)`, `i)`, `ii)`, etc.).  
            - Sublistas dentro de subitens (por exemplo, `i)`, `ii)` dentro de `a)`) também **permanecem no `content`** da cláusula-pai.  
            - Se houver uma sequência de subitens (`a)` … `p)`) após `28.2`, todos pertencem a `28.2` **até que** surja um novo marcador principal (`28.3`, `29.`, etc.).


            **Exemplo (entrada):**

              3.5. Obrigações:
                a) Segurança da informação:
                i) Criptografia ponta a ponta;
                ii) Registro de auditoria.
                b) Continuidade de negócios.
              4. Vigência.
            

            **Exemplo (saída):**
                ```json
                [
                 "clause_number": "3.5",
                 "content": "Obrigações:\na) Segurança da informação:\n   i) Criptografia ponta a ponta;\n   ii) Registro de auditoria.\nb) Continuidade de negócios."
                ]

            4. **Quebras de página e ruídos/artefatos**  
            - Ignore cabeçalhos, rodapés, numeração de página, marcas e anotações tais como:
                - ``<!-- PageHeader="..." -->``
                - ``<!-- PageFooter="..." -->``
                - ``<!-- PageBreak -->``
                - “MINUTA APROVADA”, nomes de diretoria, logotipos, números de versão, elementos de `<figure>…</figure>`.
            - Se esses elementos surgirem **dentro** da área de uma cláusula, **remova-os** do `content`.  
            - Preserve informações jurídicas úteis presentes no corpo (incluindo e-mails, prazos, números de artigos) quando fizerem parte do texto da cláusula.

            **Exemplo (entrada):**
            
             <!-- PageHeader="Companhia X" -->
                28.2. Obrigações da CONTRATADA:
                a) Tratar dados pessoais conforme instruções do controlador;
             <!-- PageFooter="V8_01_07_2025" -->


            **Exemplo (saída):**
             [
             "clause_number": "28.2",
             "content": "Obrigações da CONTRATADA:\na) Tratar dados pessoais conforme instruções do controlador;"
             ]

            
            5. **Títulos vs. cláusulas**  
            - **Não** confundir títulos/cabeçalhos com a cláusula “de verdade”.  
            - Ignore títulos por extenso ou romanos (ex.: “Cláusula Décima Quarta”, “XII”) **como título**; use-os apenas para **descobrir e normalizar `clause_number`**.  
            - Títulos como “3. DO OBJETO”, “4. RESPONSABILIDADES” **não** são cláusulas por si; a cláusula válida é o corpo iniciado em `3.1`, `3.2` etc. Se houver somente `3.` com texto corrido logo após, trate `3` como cláusula e extraia o corpo normalmente.

             **Exemplo (entrada):**
                3. DO OBJETO
                3.1. A CONTRATADA fornecerá os serviços de suporte técnico...
                3.2. A CONTRATANTE deverá disponibilizar os acessos necessários..
                            
            
             **Exemplo (saída):**
                [
                 "clause_number": "3", "content": "DO OBJETO" ,
                 "clause_number": "3.1", "content": "A CONTRATADA fornecerá os serviços de suporte técnico...",
                 "clause_number": "3.2", "content": "A CONTRATANTE deverá disponibilizar os acessos necessários..." 
            ]


            6. **Texto antes do primeiro marcador**  
            - Descarte todo texto antes do **primeiro** marcador principal do chunk.

            **Exemplo (entrada):**
                Preâmbulo, endereços e logotipos...
                1. Objeto do contrato
                A CONTRATADA...

            **Exemplo (saída):**
                [
            
                    "clause_number": "1",
                    "content": "Objeto do contrato\nA CONTRATADA..."
                ]
            

            7. **Subitens “soltos” (sem pai visível)**  
            - Se aparecerem subitens (`a)`, `b)`, `i)`) **sem** que exista um **marcador principal anterior claro** no texto disponível (página atual + até 3 seguintes), **não extraia** esses subitens.  
            - Considere-os **ambíguos** e **ignore-os**.

            8. **Ordem e fidelidade**  
            - Preserve rigorosamente a **ordem original**.  
            - **Não reescreva** o texto; apenas remova ruídos (cabeçalho/rodapé/figuras).  
            - Normalize espaços em branco excessivos e quebras de linha duplicadas.

            **Exemplo (entrada):**
                4.1. Obrigações da CONTRATADA:

                A) Manter equipe qualificada.
                B) Atender SLAs críticos.

                4.2. Multas.


            **Exemplo (saída):**

            [
                
                "clause_number": "4.1",
                "content": "Obrigações da CONTRATADA:\n\nA) Manter equipe qualificada.\nB) Atender SLAs críticos.",

                "clause_number": "4.2",
                "content": "Multas."
               
            ]
        
        **Exemplos integrados:**
          
            Exemplo integrado 1 — 1.2 com subitens a)–i):
            
              **(entrada):**

                1.2. Para todos os fins contratuais aplicáveis, a expressão “Fornecimento Contratado” ou “Fornecimento” significa, conforme o objeto disposto expressamente no item 2 do Termo de Contratação, seus Anexos ou Pedido de Compra, de forma alternativa ou conjunta:
                    a) Aquisição de Bens.
                    b) Contratação de Serviços.
                    c) Contratação de Bens e Serviços.
                    d) Contratação de Serviços de Instalação e/ou Manutenção.
                    e) Contratação de Field Service (Empreiteiras).
                    f) Contratação de Serviços de Manutenção.
                    g) Contratação de Serviços de Ativação de Dados.
                    h) Licença de Uso de Software.
                    i) Consultoria.
              **(saída):**
                 [
                    
                    "clause_number": "1.2",
                    "content": "Para todos os fins contratuais aplicáveis, a expressão “Fornecimento Contratado” ou “Fornecimento” significa, conforme o objeto disposto expressamente no item 2 do Termo de Contratação, seus Anexos ou Pedido de Compra, de forma alternativa ou conjunta:\na) Aquisição de Bens.\nb) Contratação de Serviços.\nc) Contratação de Bens e Serviços.\nd) Contratação de Serviços de Instalação e/ou Manutenção.\ne) Contratação de Field Service (Empreiteiras).\nf) Contratação de Serviços de Manutenção.\ng) Contratação de Serviços de Ativação de Dados.\nh) Licença de Uso de Software.\ni) Consultoria."
                    
                 ]
            
            Exemplo integrado 2 — 28.2 com a)–p) atravessando páginas, seguido de 28.3 e 28.4:
                **(entrada):**

                    28.2. Sem prejuízo das demais obrigações previstas nas Normas de Proteção de Dados Pessoais, a
                    CONTRATADA se obriga a:

                    a) Tratar Dados Pessoais nas hipóteses autorizadas: realizar tratamento de Dados Pessoais
                    apenas nas hipóteses autorizadas pelo controlador, conforme suas instruções, e Normas de
                    Proteção de Dados Pessoais;

                    b) Respeitar os direitos de titulares de Dados Pessoais: respeitar os direitos de titulares de dados
                    pessoais previstos nas Normas de Proteção de Dados Pessoais, conforme descrito na cláusula
                    28.4 deste instrumento;

                    ...
                    p) Transparência e prestação de contas: se e quando solicitado pela ANPD e/ou pela
                    CONTRATANTE, comprovar a adoção de medidas eficazes ao cumprimento da presente cláusula
                    e das Normas de Proteção de Dados Pessoais.

                    <!-- PageFooter="V8_01_07_2025" -->
                    <!-- PageBreak -->
                    28.3. Eliminação: os Dados Pessoais tratados no âmbito do presente Contrato serão eliminados quando
                    desnecessários, após o término de tratamento...

                    28.4. Comunicados e notificações: no âmbito do tratamento dos Dados Pessoais relacionado ao
                    presente Contrato, a CONTRATADA comunicará à CONTRATANTE por escrito, em até 48 (quarenta e
                    oito) horas, ...
                
                **(saída):**
                      "clauses": [
                        "clause_number": "28.2",
                        "content": "Sem prejuízo das demais obrigações previstas nas Normas de Proteção de Dados Pessoais, a CONTRATADA se obriga a:\na) Tratar Dados Pessoais nas hipóteses autorizadas: realizar tratamento de Dados Pessoais apenas nas hipóteses autorizadas pelo controlador, conforme suas instruções, e Normas de Proteção de Dados Pessoais;\nb) Respeitar os direitos de titulares de Dados Pessoais: respeitar os direitos de titulares de dados pessoais previstos nas Normas de Proteção de Dados Pessoais, conforme descrito na cláusula 28.4 deste instrumento;\n...\np) Transparência e prestação de contas: se e quando solicitado pela ANPD e/ou pela CONTRATANTE, comprovar a adoção de medidas eficazes ao cumprimento da presente cláusula e das Normas de Proteção de Dados Pessoais."
                        ,
                      
                        "clause_number": "28.3",
                        "content": "Eliminação: os Dados Pessoais tratados no âmbito do presente Contrato serão eliminados quando desnecessários, após o término de tratamento a que foi permitido com fundamento nas Normas de Proteção de Dados Pessoais ou quando do término deste Contrato."
                        ,
                        "clause_number": "28.4",
                        "content": "Comunicados e notificações: no âmbito do tratamento dos Dados Pessoais relacionado ao presente Contrato, a CONTRATADA comunicará à CONTRATANTE por escrito, em até 48 (quarenta e oito) horas, na hipótese da ocorrência de: (i) tratamento de Dados Pessoais não autorizado, incidente e/ou violação do disposto nesta cláusula; e/ou (ii) recebimento de notificação, reclamação, consulta ou solicitação enviada por titular de dados pessoais e/ou por uma autoridade pública, incluindo, mas não se limitando, à ANPD."
                     
                    ]
                    
        """

        # build the user message containing the actual page content
        data = f"""
        Páginas do contrato:
        {chunk["content"]}

       
        """
        # send the messages to the GPT-4o chat completion endpoint
        messages = [
            {
                "role": "system",
                "content": prompt,
            },
            {
                "role": "user",
                "content": data,
            },
        ]

        response = safe_parse_with_retry(client, messages, response_format, config)

        # update token usage metrics
        total_prompt += response.usage.prompt_tokens
        total_completion += response.usage.completion_tokens
        total_tokens += response.usage.total_tokens

        # parse the structured JSON output from the model
        structured_output = response.choices[0].message.parsed

        # store the parsed clauses under the page number (as int)
        filtered_clauses[i] = json.loads(structured_output.model_dump_json(indent=2))

        i += 1
    # return both the filtered clauses and aggregate usage statistics
    return {
        "clauses": filtered_clauses,
        "usage": {
            "prompt": total_prompt,
            "completion": total_completion,
            "total": total_tokens,
        },
    }


# %%


def filter_clauses_for_training(chunks, client, response_format, config):
    """
    Use an LLM to extract clauses, sub-clauses, and items from contract text chunks
    specifically for creating training data. Handles both numeric and letter-based
    clause numbering, preserves original numbering and content, and filters by
    Portuguese/English language preferences.

    Parameters:
        chunks (List[dict]): Each dict must contain 'page_number' and 'content'.
        client (AzureOpenAI): Initialized GPT-4o OpenAI client for chat completions.
        response_format: Desired structured output format for the LLM response.
        config (dict): Model parameters (max_tokens, temperature, top_p, deployment).

    Returns:
        dict: {
            "clauses": {page_number: PageOutput dict, ...},
            "usage": {"prompt": int, "completion": int, "total": int}
        }
    """

    # load model configuration parameters
    model_config = config

    # initialize counters for token usage
    total_prompt = 0
    total_completion = 0
    total_tokens = 0

    # container to hold filtered clauses by page
    filtered_clauses = {}

    # iterate through each page chunk
    for chunk in chunks:
        # build the system prompt with training-specific extraction rules
        prompt = """
            Você é um doutor especializado em análise de documentos jurídicos empresariais.

            Sua tarefa:

            - Extraia SOMENTE cláusulas, sub-cláusulas e subitens contratuais com base na estrutura de numeração principal (ex.: “3.”, “3.5.”, “4.1.2.”) e subitens (ex.: “a.”, “b.”, “i.”, “ii.”).
            - Alguns templates que voce vai receber a estrutura de numeração principal da cláusula será por letras (ex.: “A.”, “B.”, “C.”) e subitens por número (ex.: “A.”, “A.1.”, “A.2.”)

            REGRAS:

            1. Tanto cláusulas (formato “X.”) como sub-cláusulas (formato “X.Y.”, “X.Y.Z.”) e subitens (formato “a.”, “b.” etc.) devem ser extraídos separadamente.

            2. Se um subitem aparecer “solto” (ex.: sem número pai no mesmo chunk), ainda assim extraia-o separadamente, conforme requerido na regra 7.

            3. Preserve rigorosamente a numeração original e a sequência do texto.

            4. Se o número da cláusula vier por extenso ou em algarismos romanos (ex.: “Cláusula Décima Quarta”, “XII”), converta sempre para numeração decimal pontuada (ex.: “14.”, “12.”).

            5. Ignore o título textual (nome da cláusula) e foque apenas na numeração e no seu conteúdo.

            6. Não se esqueça de sempre seperar se a clausula ou subclausula tiver subitens, jamais junte todos os subitens em uma clausula só. Exemplo:
                    Entrada (chunk):
                    ```text
                        CLÁUSULA DÉCIMA NONA - DO TRATAMENTO DOS DADOS

                        19.5. Em caso de Incidente de Segurança envolvendo Dados Pessoais obtidos em decorrência deste Contrato, independentemente do motivo que o tenha ocasionado, deverá a CONTRATADA, imediatamente, apresentar a CLARO relato com, no mínimo, as seguintes informações:
                        a) data e hora do Incidente de Segurança;
                        b) data e hora da ciência do Incidente pela CONTRATADA;
                    Saída:
                        "page_number":  X,
                        "clauses": [
                            "clause_title": DO TRATAMENTO DOS DADOS
                            "clause_number": "19.5.a..",
                            "content": "data e hora do Incidente de Segurança;"
                            ,

                            "clause_number": "19.5.b.",
                            "content": "data e hora da ciência do Incidente pela CONTRATADA;"
                            ,
                            …
                        ]  
                    **- Se o chunk só tiver subitens listado na página, sem a possibilidade de saber a clásula pai, extraia somente o subitem normalmente como clause_number. Exemplo:**
                    Entrada (chunk):
                    ```text
                        g. deverá manter registros completos e precisos e documentos de suporte adequados em relação à sua prestação de serviços e fornecerá à EMPRESA e/ou ao seu representante autorizado, acesso total a esses registros, documentos de suporte e informações necessárias para demonstrar a conformidade com as leis de privacidade e de proteção de dados aplicável;
                        h. não recebeu quaisquer solicitações por parte de qualquer autoridade ou agência para qualquer transferência de Dados Pessoais fornecidos a ele ou em nome da EMPRESA ou acessadas de outra forma em relação à prestação de serviços, nem está ciente de qualquer solicitação não atendida advinda de qualquer autoridade ou agência para qualquer transferência de tais Dados Pessoais;
                    Saída:
                        "page_number": X,
                            "clauses": [
                                "clause_title": DO TRATAMENTO DOS DADOS
                                "clause_number": "g",
                                "content": "deverá manter registros completos e precisos e documentos…"
                                ,
                                "clause_title": DO TRATAMENTO DOS DADOS
                                "clause_number": "h.",
                                "content":"não recebeu quaisquer solicitações por…"
                                ,
                                …
                            ]
                7. Se a pagina tiver as cláusulas em portugues e inglês, sempre pegue a parte em português. Se tiver somente em inglês, extraia em inglês.
        """
        # build the user message containing the page content
        data = f"""
        Página do contrato:
        {chunk["content"]}


        **Por favor não ignore subitens soltos mesmo que não seja possivel identificar a clausula pai, liste-os normalmente como "clause_number".
        **Ignore titulos de cláusulas e não confunda-os como uma clausula em si conforme requerido na regra 4."

       
        """
        # send the prompt and page content to the LLM for parsing

        messages = [
            {
                "role": "system",
                "content": prompt,
            },
            {
                "role": "user",
                "content": data,
            },
        ]

        response = safe_parse_with_retry(client, messages, response_format, config)

        # update token usage counters after the call
        total_prompt += response.usage.prompt_tokens
        total_completion += response.usage.completion_tokens
        total_tokens += response.usage.total_tokens

        # parse the structured JSON output from the LLM
        structured_output = response.choices[0].message.parsed

        # store the parsed PageOutput JSON under the integer page number
        filtered_clauses[int(chunk["page_number"])] = json.loads(
            structured_output.model_dump_json(indent=2)
        )

    # return the filtered clauses along with aggregated usage stats
    return {
        "clauses": filtered_clauses,
        "usage": {
            "prompt": total_prompt,
            "completion": total_completion,
            "total": total_tokens,
        },
    }


# %%
def create_final_document_with_bubbles(reviewed_data, output_docx):
    """
    Generate a Word document highlighting original vs. revised clauses with change bubbles.

    This function takes the structured review results, creates a .docx where:
      - Each clause’s original text is shown with strikethrough in red.
      - Each clause’s revised text follows, underlined in green.
      - A comment bubble is added to each clause paragraph containing the legal issue.

    The document is first built in a temporary file, comments are inserted, then it is
    moved to the final output path and cleaned up.

    Parameters:
        reviewed_data (dict): Mapping of any keys to dicts with a 'clauses' list.
            Each clause dict should include:
              - 'numero_da_clausula' (str)
              - either 'clasula_original' or 'original' (str)
              - either 'clausula_revisada' or 'revised' (str)
              - 'problema_juridico' (str) for the comment bubble text
        output_docx (str): Filesystem path where the final .docx will be saved.
    """

    # Create a new in-memory Word document
    doc = Document()

    # First pass: add each clause with styling (strike/red for original, underline/green for revised)
    for page in reviewed_data.values():
        for clause in page["clauses"]:
            # fetch clause number or default to empty
            numero = clause.get("numero_da_clausula", "")
            # get original text, fallback to 'original' key
            original_text = clause.get("clasula_original", clause.get("original", ""))
            # get revised text, fallback to 'revised' key
            revised_text = clause.get("clausula_revisada", clause.get("revised", ""))

            # Combine number and original text for display
            full_orig = f"{numero} – {original_text}"
            # Add a new paragraph for this clause
            p = doc.add_paragraph()

            # Add the original text run, styled with strikethrough and red color
            run_o = p.add_run(full_orig + " ")
            run_o.font.strike = True
            run_o.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

            # Add the revised text run, underlined and colored green
            run_r = p.add_run(revised_text)
            run_r.font.underline = True
            run_r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    # Create a temporary directory to hold the intermediary document
    tmp_dir = tempfile.mkdtemp()
    tmp_path = Path(tmp_dir) / "diff.docx"
    # Save the styled document to the temporary path
    doc.save(str(tmp_path))

    # Second pass: add comment bubbles for each clause paragraph
    para_index = 0
    for page in reviewed_data.values():
        for clause in page["clauses"]:
            # The comment text is the identified legal issue
            comment = clause.get("problema_juridico", "")
            add_comment_bubble_opc(
                input_path=str(tmp_path),  # read/write the same temp file
                output_path=str(tmp_path),  # apply bubble to this file
                para_index=para_index,  # which paragraph to attach comment to
                comment_text=comment,  # content of the comment bubble
            )
            para_index += 1  # move to the next paragraph index

    # Ensure the output directory exists
    Path(output_docx).parent.mkdir(exist_ok=True, parents=True)
    # Move the finalized document from temp to the desired output path
    shutil.move(str(tmp_path), output_docx)
    # Clean up the temporary directory
    shutil.rmtree(tmp_dir)
    # Inform the user of successful save
    print(f"Documento simples com diffs e bubbles salvo em {output_docx}")


# %%


def create_original_and_revised_docs(
    reviewed_data: dict, out_dir: Path, blob_name: str
):
    """
    Generate separate Word documents for original and revised clauses, adding comment bubbles to the revised version.

    This function:
      - Derives a base filename from `blob_name`.
      - Creates two .docx documents:
          1. `{base}-original.docx` containing each clause’s original text.
          2. `{base}-revisado.docx` containing each clause’s revised text with comment bubbles.
      - Saves both files under `out_dir`, creating the directory if needed.
      - Returns the paths to the original and revised documents.

    Parameters:
        reviewed_data (dict):
            Mapping of pages to dicts with a `'clauses'` list. Each clause dict must include:
              - `'numero_da_clausula'` (str): Clause identifier.
              - `'clasula_original'` (str): Original clause text.
              - `'clausula_revisada'` (str): Revised clause text.
              - `'problema_juridico'` (str): Text for the comment bubble.
        out_dir (Path):
            Directory where output documents will be created.
        blob_name (str):
            Original filename (used to derive the base name for output files).

    Returns:
        Tuple[Path, Path]: Paths to the original document and the final revised document.
    """

    # Derive the base name (without extension) from the blob name
    base = Path(blob_name).name.split(".", 1)[0]

    # Initialize Word documents for original and revised clauses
    orig = Document()

    rev = Document()

    # Collect all comment texts to add later
    clauses = []
    for page in reviewed_data.values():
        for clause in page["clauses"]:
            # Extract fields from the reviewed clause
            num = clause.get("numero_da_clausula", "")
            orig_text = clause.get("clasula_original", "")
            rev_text = clause.get("clausula_revisada", "")
            # comment   = clause.get('problema_juridico', '')

            # Add a paragraph for the original text
            orig.add_paragraph(f"{num} - {orig_text}")
            # Add a paragraph for the revised text
            rev.add_paragraph(f"{num} - {rev_text}")
            # Store the comment for bubble insertion
            # clauses.append(comment)

    # Ensure the output directory exists
    out_dir.mkdir(parents=True, exist_ok=True)
    # Define file paths for original, temporary revised, and final revised docs
    orig_path = out_dir / f"{base}-original.docx"
    tmp_rev_path = out_dir / f"{base}-revisado_tmp.docx"
    final_rev_path = out_dir / f"{base}-revisado.docx"
    # Save the original and temporary revised documents
    orig.save(str(orig_path))
    rev.save(str(tmp_rev_path))

    # Insert comment bubbles into each paragraph of the revised document
    # for idx, comment in enumerate(clauses):
    # if comment:
    # add_comment_bubble_opc(
    # input_path=str(tmp_rev_path),   # read/write the same temp file
    # output_path=str(tmp_rev_path),  # overwrite it with comments added
    # para_index=idx,             # paragraph index to annotate
    # comment_text=comment    # the bubble text
    # )

    # Remove any existing final revised doc, then rename the temp file
    if final_rev_path.exists():
        final_rev_path.unlink()
    tmp_rev_path.rename(final_rev_path)

    # Return the paths to the created documents
    return orig_path, final_rev_path


# %%


def add_comment_bubble_opc(
    input_path,
    output_path,
    para_index,
    comment_text,
    author="LLM-Review",
    initials="LR",
):
    """
    Insert a comment bubble into a .docx (OPC) Word document at the specified paragraph index.

    This function:
      1. Unzips the .docx file to a temporary folder.
      2. Ensures comments are enabled in [Content_Types].xml.
      3. Loads or creates word/comments.xml and relationships.
      4. Appends a new <w:comment> element with the given text, author, and initials.
      5. Inserts commentRangeStart/End and a commentReference into the target paragraph.
      6. Rezips the folder back into a .docx at `output_path`.
      7. Cleans up temporary files and prints a confirmation message.

    Parameters:
        input_path (str):  Path to the source .docx file.
        output_path (str): Path where the modified .docx will be written.
        para_index (int):  Zero-based index of the paragraph to annotate.
        comment_text (str): The text to display in the comment bubble.
        author (str):      Author name for the comment metadata (default 'LLM-Review').
        initials (str):    Initials for the comment metadata (default 'LR').
    """

    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(input_path, "r") as zin:
            zin.extractall(temp_dir)

        # Define XML namespaces for WordprocessingML and OPC
        NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        NS_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
        NS_CT = "http://schemas.openxmlformats.org/package/2006/content-types"
        NSMAP_W = {"w": NS_W}

        ct_path = os.path.join(temp_dir, "[Content_Types].xml")
        ct_tree = etree.parse(ct_path)
        ct_root = ct_tree.getroot()
        exists = ct_root.xpath(
            "ct:Override[@PartName='/word/comments.xml']", namespaces={"ct": NS_CT}
        )
        if not exists:
            override = etree.SubElement(ct_root, f"{{{NS_CT}}}Override")
            override.set("PartName", "/word/comments.xml")
            override.set(
                "ContentType",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
            )
            ct_tree.write(
                ct_path, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        comments_path = os.path.join(temp_dir, "word", "comments.xml")
        if os.path.exists(comments_path):
            tree_c = etree.parse(comments_path)
            root_c = tree_c.getroot()
            existing_ids = [
                int(c.get(f"{{{NS_W}}}id"))
                for c in root_c.findall("w:comment", NSMAP_W)
                if c.get(f"{{{NS_W}}}id") is not None
            ]
        else:
            root_c = etree.Element(f"{{{NS_W}}}comments", nsmap=NSMAP_W)
            tree_c = etree.ElementTree(root_c)
            existing_ids = []

            rels_path = os.path.join(temp_dir, "word", "_rels", "document.xml.rels")
            rels_tree = etree.parse(rels_path)
            rels_root = rels_tree.getroot()
            etree.SubElement(
                rels_root,
                f"{{{NS_REL}}}Relationship",
                {
                    "Id": f"rId{len(rels_root) + 1}",
                    "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
                    "Target": "comments.xml",
                },
            )
            rels_tree.write(
                rels_path, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        new_id = max(existing_ids) + 1 if existing_ids else 0

        comment_elem = etree.Element(f"{{{NS_W}}}comment", nsmap=NSMAP_W)
        comment_elem.set(f"{{{NS_W}}}id", str(new_id))
        comment_elem.set(f"{{{NS_W}}}author", author)
        comment_elem.set(f"{{{NS_W}}}initials", initials)
        comment_elem.set(f"{{{NS_W}}}date", "")
        p = etree.SubElement(comment_elem, f"{{{NS_W}}}p")
        r = etree.SubElement(p, f"{{{NS_W}}}r")
        t = etree.SubElement(r, f"{{{NS_W}}}t")
        t.text = comment_text
        root_c.append(comment_elem)

        tree_c.write(
            comments_path, xml_declaration=True, encoding="UTF-8", standalone=True
        )

        doc_path = os.path.join(temp_dir, "word", "document.xml")
        tree_d = etree.parse(doc_path)
        root_d = tree_d.getroot()
        paras = root_d.findall(".//w:p", NSMAP_W)
        if para_index < len(paras):
            target = paras[para_index]

            crs = etree.Element(f"{{{NS_W}}}commentRangeStart", nsmap=NSMAP_W)
            crs.set(f"{{{NS_W}}}id", str(new_id))
            target.insert(0, crs)

            cre = etree.Element(f"{{{NS_W}}}commentRangeEnd", nsmap=NSMAP_W)
            cre.set(f"{{{NS_W}}}id", str(new_id))
            target.append(cre)

            ref_run = etree.SubElement(target, f"{{{NS_W}}}r")
            cref = etree.SubElement(
                ref_run, f"{{{NS_W}}}commentReference", nsmap=NSMAP_W
            )
            cref.set(f"{{{NS_W}}}id", str(new_id))
        tree_d.write(doc_path, xml_declaration=True, encoding="UTF-8", standalone=True)

        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    full_path = os.path.join(root, file)
                    arcname = os.path.relpath(full_path, temp_dir)
                    zout.write(full_path, arcname)

    finally:
        shutil.rmtree(temp_dir)
        print(
            f"Comentario OPC adicionado no paragrafo {para_index}, salvo em {output_path}"
        )


def create_temp_index(search_index, deployment_name, model_name, resource_url):
    """
    Create a temporary Azure Cognitive Search index for clause embeddings.

    Parameters:
        search_index: SearchIndexClient for managing indexes.
        deployment_name (str): Azure OpenAI deployment for the vectorizer.
        model_name (str): Embedding model name (e.g., "text-embedding-ada-002").
        resource_url (str): Azure OpenAI resource URL (from Settings).

    Returns:
        str: The name of the newly created temporary index.
    """
    tmp_index_name = f"tmp_clause_{uuid.uuid4().hex[:8]}"

    fields = [
        SearchField(name="id", type=SearchFieldDataType.String, key=True),
        SearchField(
            name="numero_da_clausula",
            type=SearchFieldDataType.String,
            filterable=True,
            searchable=True,
        ),
        SearchField(
            name="clasula_original", type=SearchFieldDataType.String, searchable=True
        ),
        SearchField(
            name="problema_juridico", type=SearchFieldDataType.String, searchable=True
        ),
        SearchField(
            name="clausula_revisada", type=SearchFieldDataType.String, searchable=True
        ),
        SearchField(
            name="embedding",
            type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
            vector_search_dimensions=1536,
            vector_search_profile_name="myHnswProfile",
        ),
    ]

    vector_search = VectorSearch(
        algorithms=[HnswAlgorithmConfiguration(name="myHnsw")],
        profiles=[
            VectorSearchProfile(
                name="myHnswProfile",
                algorithm_configuration_name="myHnsw",
                vectorizer_name="clauseVector",
            )
        ],
        vectorizers=[
            AzureOpenAIVectorizer(
                vectorizer_name="clauseVector",
                kind="azureOpenAI",
                parameters=AzureOpenAIVectorizerParameters(
                    resource_url=resource_url,
                    deployment_name=deployment_name,
                    auth_identity=SearchIndexerDataNoneIdentity(),
                    model_name=model_name,
                ),
            )
        ],
    )

    index = SearchIndex(name=tmp_index_name, fields=fields, vector_search=vector_search)
    result = search_index.create_or_update_index(index)
    return tmp_index_name


def vectorize_and_upload(data, index_client, embeddings_client):
    """
    Generate embeddings for each contract clause and upload to Azure Cognitive Search.

    Parameters:
        data (dict): Mapping of page keys to dicts with 'clauses' lists.
        index_client: SearchClient with upload_documents method.
        embeddings_client: AzureOpenAI client configured for embeddings.
    """
    documents = []
    for page_key, page in data.items():
        for clause in page["clauses"]:
            documents.append(
                {
                    "id": clause["id"],
                    "numero_da_clausula": clause["numero_da_clausula"],
                    "clasula_original": clause["clasula_original"],
                    "problema_juridico": clause["problema_juridico"],
                    "clausula_revisada": clause["clausula_revisada"],
                    "embedding": generate_embedding(
                        embeddings_client, clause["clasula_original"]
                    ),
                }
            )
    result = index_client.upload_documents(documents=documents)
    return len(result)


# Parâmetros de retry — ajuste conforme sua realidade
MAX_RETRIES = 3  # número de tentativas por chunk
INITIAL_BACKOFF = 5  # segundos iniciais de espera
REQUEST_TIMEOUT = None  # já controlado no http_client do AzureOpenAI


def safe_parse_with_retry(client, messages, response_format, model_config):
    """
    Envia mensagens ao LLM aplicando retry exponencial em HttpResponseError.
    Retorna o objeto resposta em caso de sucesso, ou levanta a última exceção.
    """
    backoff = INITIAL_BACKOFF
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            return client.beta.chat.completions.parse(
                messages=messages,
                response_format=response_format,
                max_tokens=model_config["max_tokens"],
                temperature=model_config["temperature"],
                top_p=model_config["top_p"],
                model=model_config["deployment"],
            )
        except HttpResponseError as e:
            code = getattr(e, "status_code", None)
            print(
                f"[LLM Retry] tentativa {attempt}/{MAX_RETRIES} falhou "
                f"(status={code}): {e}"
            )
            if attempt == MAX_RETRIES:
                print("Exauridas as tentativas de retry no LLM.")
                raise
            print(f"Aguardando {backoff}s antes da próxima tentativa…")
            time.sleep(backoff)
            backoff *= 2


def deduplicate_clauses(extracted_pages: dict) -> list[dict]:
    """
    Mescla cláusulas sobrepostas vindas de múltiplas páginas em uma lista única.

    Parameters:
        extracted_pages (dict): dicionário com a chave 'clauses', cujo valor é outro dict:
            {
                '0': {'page_number': int, 'clauses': [ {clause_number, content}, ... ]},
                '1': {...},
                ...
            }

    Returns:
        list[dict]: lista de cláusulas únicas, cada uma com 'clause_number' e 'content',
                    na ordem em que cada número apareceu pela primeira vez.
    """

    def normalize_spaces(text: str) -> str:
        # remove quebras de linha e espaços múltiplos
        return " ".join(text.split())

    # 1) Extrai e ordena por page_number
    raw_chunks = extracted_pages.get("clauses", {})
    page_list = list(raw_chunks.values())
    page_list.sort(key=lambda chunk: chunk.get("page_number", 0))

    consolidated: dict[str, str] = {}
    clause_order: list[str] = []

    # 2) Processa cada cláusula em cada chunk
    for page_data in page_list:
        for clause in page_data.get("clauses", []):
            key = normalize_clause_number(clause["clause_number"])
            raw_content = clause["content"].strip()
            norm_content = normalize_spaces(raw_content)

            if key not in consolidated:
                # primeira vez que vemos este número
                consolidated[key] = raw_content
                clause_order.append(key)
            else:
                # já vimos: substitui se o novo for "mais longo" (mais informação)
                existing_norm = normalize_spaces(consolidated[key])
                if len(norm_content) > len(existing_norm):
                    consolidated[key] = raw_content

    # 3) Monta lista final mantendo a ordem de aparição
    final_clauses = [
        {"clause_number": num, "content": consolidated[num]} for num in clause_order
    ]
    return final_clauses
