"""
Word document generation for reviewed contracts.

Creates .docx files with original and revised clauses,
including OPC comment bubble annotations for legal issues.
"""

import os
import logging
import tempfile
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import RGBColor
from lxml import etree

logger = logging.getLogger(__name__)


def create_original_and_revised_docs(
    reviewed_data: dict, out_dir: Path, blob_name: str
):
    """
    Generate separate Word documents for original and revised clauses.

    Parameters:
        reviewed_data (dict): Mapping of pages to dicts with 'clauses' lists.
        out_dir (Path): Directory for output documents.
        blob_name (str): Original filename (used for output naming).

    Returns:
        tuple[Path, Path]: Paths to the original and revised documents.
    """
    base = Path(blob_name).name.split(".", 1)[0]

    orig = Document()
    rev = Document()

    for page in reviewed_data.values():
        for clause in page["clauses"]:
            num = clause.get("numero_da_clausula", "")
            orig_text = clause.get("clasula_original", "")
            rev_text = clause.get("clausula_revisada", "")

            orig.add_paragraph(f"{num} - {orig_text}")
            rev.add_paragraph(f"{num} - {rev_text}")

    out_dir.mkdir(parents=True, exist_ok=True)
    orig_path = out_dir / f"{base}-original.docx"
    tmp_rev_path = out_dir / f"{base}-revisado_tmp.docx"
    final_rev_path = out_dir / f"{base}-revisado.docx"

    orig.save(str(orig_path))
    rev.save(str(tmp_rev_path))

    if final_rev_path.exists():
        final_rev_path.unlink()
    tmp_rev_path.rename(final_rev_path)

    return orig_path, final_rev_path


def create_final_document_with_bubbles(reviewed_data, output_docx):
    """
    Generate a Word document with red strikethrough (original) and green underline
    (revised) for each clause, plus comment bubbles with legal issues.

    Parameters:
        reviewed_data (dict): Mapping of pages to dicts with 'clauses' lists.
        output_docx (str): Path where the final .docx will be saved.
    """
    doc = Document()

    for page in reviewed_data.values():
        for clause in page["clauses"]:
            numero = clause.get("numero_da_clausula", "")
            original_text = clause.get("clasula_original", clause.get("original", ""))
            revised_text = clause.get("clausula_revisada", clause.get("revised", ""))

            full_orig = f"{numero} – {original_text}"
            p = doc.add_paragraph()

            run_o = p.add_run(full_orig + " ")
            run_o.font.strike = True
            run_o.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

            run_r = p.add_run(revised_text)
            run_r.font.underline = True
            run_r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    tmp_dir = tempfile.mkdtemp()
    tmp_path = Path(tmp_dir) / "diff.docx"
    doc.save(str(tmp_path))

    para_index = 0
    for page in reviewed_data.values():
        for clause in page["clauses"]:
            comment = clause.get("problema_juridico", "")
            add_comment_bubble_opc(
                input_path=str(tmp_path),
                output_path=str(tmp_path),
                para_index=para_index,
                comment_text=comment,
            )
            para_index += 1

    Path(output_docx).parent.mkdir(exist_ok=True, parents=True)
    shutil.move(str(tmp_path), output_docx)
    shutil.rmtree(tmp_dir)
    logger.info("Document with diffs and bubbles saved to %s", output_docx)


def add_comment_bubble_opc(
    input_path,
    output_path,
    para_index,
    comment_text,
    author="LLM-Review",
    initials="LR",
):
    """
    Insert a comment bubble into a .docx at the specified paragraph index.

    Manipulates the OPC (Open Packaging Conventions) XML directly:
    unzips the .docx, adds/modifies comments.xml, inserts comment
    range markers into the target paragraph, and rezips.

    Parameters:
        input_path (str): Path to the source .docx file.
        output_path (str): Path where the modified .docx will be written.
        para_index (int): Zero-based paragraph index to annotate.
        comment_text (str): Text for the comment bubble.
        author (str): Author name for comment metadata.
        initials (str): Initials for comment metadata.
    """
    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(input_path, "r") as zin:
            zin.extractall(temp_dir)

        NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        NS_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
        NS_CT = "http://schemas.openxmlformats.org/package/2006/content-types"
        NSMAP_W = {"w": NS_W}

        ct_path = os.path.join(temp_dir, "[Content_Types].xml")
        ct_tree = etree.parse(ct_path)
        ct_root = ct_tree.getroot()
        exists = ct_root.xpath(
            "ct:Override[@PartName='/word/comments.xml']", namespaces={"ct": NS_CT}
        )
        if not exists:
            override = etree.SubElement(ct_root, f"{{{NS_CT}}}Override")
            override.set("PartName", "/word/comments.xml")
            override.set(
                "ContentType",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
            )
            ct_tree.write(
                ct_path, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        comments_path = os.path.join(temp_dir, "word", "comments.xml")
        if os.path.exists(comments_path):
            tree_c = etree.parse(comments_path)
            root_c = tree_c.getroot()
            existing_ids = [
                int(c.get(f"{{{NS_W}}}id"))
                for c in root_c.findall("w:comment", NSMAP_W)
                if c.get(f"{{{NS_W}}}id") is not None
            ]
        else:
            root_c = etree.Element(f"{{{NS_W}}}comments", nsmap=NSMAP_W)
            tree_c = etree.ElementTree(root_c)
            existing_ids = []

            rels_path = os.path.join(temp_dir, "word", "_rels", "document.xml.rels")
            rels_tree = etree.parse(rels_path)
            rels_root = rels_tree.getroot()
            etree.SubElement(
                rels_root,
                f"{{{NS_REL}}}Relationship",
                {
                    "Id": f"rId{len(rels_root) + 1}",
                    "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
                    "Target": "comments.xml",
                },
            )
            rels_tree.write(
                rels_path, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        new_id = max(existing_ids) + 1 if existing_ids else 0

        comment_elem = etree.Element(f"{{{NS_W}}}comment", nsmap=NSMAP_W)
        comment_elem.set(f"{{{NS_W}}}id", str(new_id))
        comment_elem.set(f"{{{NS_W}}}author", author)
        comment_elem.set(f"{{{NS_W}}}initials", initials)
        comment_elem.set(f"{{{NS_W}}}date", "")
        p = etree.SubElement(comment_elem, f"{{{NS_W}}}p")
        r = etree.SubElement(p, f"{{{NS_W}}}r")
        t = etree.SubElement(r, f"{{{NS_W}}}t")
        t.text = comment_text
        root_c.append(comment_elem)

        tree_c.write(
            comments_path, xml_declaration=True, encoding="UTF-8", standalone=True
        )

        doc_path = os.path.join(temp_dir, "word", "document.xml")
        tree_d = etree.parse(doc_path)
        root_d = tree_d.getroot()
        paras = root_d.findall(".//w:p", NSMAP_W)
        if para_index < len(paras):
            target = paras[para_index]

            crs = etree.Element(f"{{{NS_W}}}commentRangeStart", nsmap=NSMAP_W)
            crs.set(f"{{{NS_W}}}id", str(new_id))
            target.insert(0, crs)

            cre = etree.Element(f"{{{NS_W}}}commentRangeEnd", nsmap=NSMAP_W)
            cre.set(f"{{{NS_W}}}id", str(new_id))
            target.append(cre)

            ref_run = etree.SubElement(target, f"{{{NS_W}}}r")
            cref = etree.SubElement(
                ref_run, f"{{{NS_W}}}commentReference", nsmap=NSMAP_W
            )
            cref.set(f"{{{NS_W}}}id", str(new_id))
        tree_d.write(doc_path, xml_declaration=True, encoding="UTF-8", standalone=True)

        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    full_path = os.path.join(root, file)
                    arcname = os.path.relpath(full_path, temp_dir)
                    zout.write(full_path, arcname)

    finally:
        shutil.rmtree(temp_dir)
        logger.info("Comment added at paragraph %d, saved to %s", para_index, output_path)
